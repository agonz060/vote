#!/usr/local/bin/python2.7

# import module docx creator and mysql connection modules
import cgi
import sys
import mysql.connector
from docx import Document
from docx.shared import Inches
from pythonMysqlConfig import readDbConfig
from createDocxFunctions import getWorkingDir, getResultsFilePath
from createDocxFunctions import intToRoman, mergeColumns, setColumnWidth, printComments
from createDocxFunctions import getTodaysDate, getPollData, getVoteCounts, formatActionResults
from createDocxFunctions import getActionDescriptions, isMultiActionPollType, checkForEvaluations

# variables
fileName = poll_id = None

# testing
# poll_id = 16 # poll type: fifth year review
# poll_id = 17 # poll type: fifth year appraisal
# poll_id = 18 # poll type: merit
# fileName = 'test.docx'

# load document settings
if(len(sys.argv) == 3):
	poll_id = sys.argv[1]
  	# fileName = '/usr/local/www/engr/cmsengr/www-root/intranet/vote/admin/createDocx/tmp/{}'.format(sys.argv[2])
  	fileName = sys.argv[2]
  	# print('poll_id: {0} fileName: {1}'.format(poll_id,fileName))
  	# fileName
else:
	print(cgi.escape('createDocx usage: ./createDocx.py <poll_id> <saveFileAsName>'))
	sys.exit(1)

# connect to database, print error if any
dbConfig = readDbConfig()
try:
	conn = mysql.connector.connect(**dbConfig)
except mysql.connector.Error as err:
	print(err)
	sys.exit(1)

# constants
TRUE = 1
FALSE = 0
COL_ONE = 0
COL_TWO = 1
COL_THREE = 2
IS_EVALUATION = TRUE
TEMPLATE = 'template.docx'

# various notices/headings
VOTES_HEADING = 'Votes - (to be recorded on Dept. Letter)'
COMMENTS_NOTICE = '(Comments are transcribed from ballots as written. Whether comments were discussed at the meeting'
COMMENTS_NOTICE += ' and therefore included in the department letter is up to the chair to determine.)'
EVALUATIONS_HEADING = 'Confidential Evaluations - (to be recorded on Dept. Letter)'
TOTAL_ELIGIBLE_NOTICE = 'Total Eligible voting members: '

# create dictionary for easier access/modification
HEADINGS = { 'VOTES' : VOTES_HEADING, 'EVALUATIONS' : EVALUATIONS_HEADING }
NOTICES = { 'COMMENTS' : COMMENTS_NOTICE , 'ELIGIBLE' : TOTAL_ELIGIBLE_NOTICE }

# print('0')
# open document, this document contains a custom borderless table styles
workingDir = tempFilePath = templateFile = None

# form path to template.docx file
workingDir = getWorkingDir()
tempFilePath = workingDir + TEMPLATE
# print('templateFilePath: {}'.format(tempFilePath))
try:
	templateFile = open(tempFilePath,'rb')
except:
	print ('Error: couldn\'t open template.docx')
	sys.exit(1)
# print('1')
try:
	doc = Document(templateFile)
except:
	print('Error: couldn\'t create docx file from template.docx')
	sys.exit(1)

# print('2')
# create table for displaying result heading
headingTable = doc.add_table(rows=2, cols=2)

# change table style to custom table style that hides border lines
headingTable.style = 'borderlessTable'
setColumnWidth(headingTable.columns[COL_ONE],Inches(1))
setColumnWidth(headingTable.columns[COL_TWO],Inches(5.5))

# add date information to first row of the heading table
headingTable.cell(0,0).text = 'DATE:'
headingTable.cell(0,1).text = getTodaysDate()

# add description second row
try:
	# get description from poll data
	pollData = getPollData(conn,poll_id)
	descriptions = getActionDescriptions(pollData)

	# display discription
	headingTable.cell(1,0).text = 'RE:'
	headingDescription = descriptions[0]
	headingTable.cell(1,1).text = headingDescription

except:
	print "Could not retreive poll description"
	sys.exit(1)

# format horizontal line (divider)
doc.add_paragraph('', style='horizontalLine')
doc.add_paragraph('', style='emptyLine')

# poll variables
# pollData and descriptions set on line 80
voteCounts = getVoteCounts(conn,pollData)
pollType = pollData['pollType']

# # printing variables
actionNum = eligible = 0

for description in descriptions:
	# begin from first action, subtract 1 from actionNum for indexing
	actionNum += 1
	romanActionNum = intToRoman(actionNum)

	# print action vote tally before printing action evaluations (comments)
	printActionVoteResults = TRUE
	printActionEvaluations = FALSE

	# print both types of action results/information
	while(printActionVoteResults or printActionEvaluations):
		# reset variables
		eligible = 0
		# print('2.2')
		# start with vote tally
		if(printActionVoteResults):
			# set action result variables
			if(isMultiActionPollType(pollType)):
				eligible = voteCounts[str(actionNum)]['eligible']
			else:
				eligible = voteCounts['eligible']

			# print('2.4')
			actionResults = formatActionResults(conn,pollData,actionNum)
			eligibleSummary = NOTICES['ELIGIBLE'] + str(eligible)
			# print('2.4.5')
			# create table to format results
			actionTable = doc.add_table(rows=5,cols=3)
			actionTable.style = 'borderlessTable'
			# print('2.5')
			setColumnWidth(actionTable.columns[COL_ONE],Inches(0.5))
			setColumnWidth(actionTable.columns[COL_TWO],Inches(1.2))
			setColumnWidth(actionTable.columns[COL_THREE],Inches(4.8))
			mergeColumns(COL_TWO,COL_THREE,actionTable)
			# print('2.6')
			# display action num and votes heading
			actionTable.cell(0,0).text = romanActionNum
			# print heading
			actionTable.cell(0,1).text = HEADINGS['VOTES']
			# print action description
			actionTable.cell(1,1).text = description
			# print total eligible notice
			actionTable.cell(2,1).text = eligibleSummary
			# print action results
			actionTable.cell(3,1).text = actionResults
			# print comments notice
			actionTable.cell(4,1).text = NOTICES['COMMENTS']

			# begin printing action comments, if any
			printComments(actionTable,conn,pollData,actionNum)
			# print('2.8')
			# format horizontal line (divider)
			doc.add_paragraph('', style='horizontalLine')
			doc.add_paragraph('', style='emptyLine')

			# end of printing vote tally and comments, start printing evaluations
			printActionVoteResults = FALSE
			printActionEvaluations = TRUE
			# print('2.9')
		elif(printActionEvaluations):
			# print('3.2')
			# there may not be evaluations for this action
			evaluations = checkForEvaluations(conn,pollData,actionNum)

			# print evaluations if there are comments
			if(evaluations['commentCount'] > 0):
				# evaluation variables
				eligibleSummary = NOTICES['ELIGIBLE'] + str(evaluations['eligible'])

				# create table to format evaluations
				evaluationsTable = doc.add_table(rows=4,cols=3)
				evaluationsTable.style = 'borderlessTable'
				setColumnWidth(evaluationsTable.columns[COL_ONE],Inches(0.5))
				setColumnWidth(evaluationsTable.columns[COL_TWO],Inches(1.2))
				setColumnWidth(evaluationsTable.columns[COL_THREE],Inches(4.8))
				mergeColumns(COL_TWO,COL_THREE,evaluationsTable)
				
				# display action num and votes heading
				evaluationsTable.cell(0,0).text = romanActionNum + 'a'
				# print heading
				evaluationsTable.cell(0,1).text = HEADINGS['EVALUATIONS']
				# print action description
				evaluationsTable.cell(1,1).text = description
				# print total eligible notice
				evaluationsTable.cell(2,1).text = eligibleSummary
				# print comments notice
				evaluationsTable.cell(3,1).text = NOTICES['COMMENTS']
				# print comments
				printComments(evaluationsTable,conn,pollData,actionNum,IS_EVALUATION)

				# format horizontal line (divider)
				doc.add_paragraph('', style='horizontalLine')
				doc.add_paragraph('', style='emptyLine')

			# begin printing action comments, if any
			# end of printing evaluations
			printActionEvaluations = FALSE

# close database connection
conn.close()

# print('3')
# setup results file page and create file for writing docx file to 
resultsFile = None
resultsFilePath = getResultsFilePath() + fileName
# print('fileName: {}'.format(fileName))
# print('Results file path: {}'.format(resultsFilePath))
try:
	resultsFile = open(resultsFilePath,'wb')

	# create docx file, then close results file
	doc.save(resultsFile)
	resultsFile.close()
except:
	print('Error - couldn\'t open results file: {}'.format(resultsFilePath))
	sys.exit(1)

# print('4')

# close template file
templateFile.close()

# print success message
print('Success')
sys.exit(0)