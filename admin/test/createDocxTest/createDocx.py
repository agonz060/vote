#!/usr/local/bin/python2.7

# import module docx creator and mysql connection modules
import sys
import mysql.connector
from docx import Document
from docx.shared import Inches
# from datetime import datetime
from pythonMysqlConfig import readDbConfig
from createDocxFunctions import intToRoman, setColumnWidth, printComments
from createDocxFunctions import getTodaysDate, getPollData, getVoteCounts, formatActionResults
from createDocxFunctions import getActionDescriptions, isMultiActionPollType, checkForEvaluations


# get poll id to begin loading poll data
pollId = 16 # poll type: fifth year review
# pollId = 17 # poll type: fifth year appraisal
# pollId = 18 # poll type: merit
# iterate through arguments
# if(len(sys.argv) == 2):
# 	pollId = sys.argv[1]
# else:
# 	print "createDocx: error - poll id required"

# connect to database, print error if any
dbConfig = readDbConfig()
try:
	conn = mysql.connector.connect(**dbConfig)
except mysql.connector.Error as err:
	print(err)

# constants
TRUE = 1
FALSE = 0
COL_ONE = 0
COL_TWO = 1
COL_THREE = 2
IS_EVALUATION = TRUE

# various notices/headings
VOTES_HEADING = 'Votes - (to be recorded on Dept. Letter)'
COMMENTS_NOTICE = '(Comments are transcribed from ballots as written. Whether comments were discussed at the meeting'
COMMENTS_NOTICE += ' and therefore included in the department letter is up to the chair to determine.)'
EVALUATIONS_HEADING = 'Confidential Evaluations - (to be recorded on Dept. Letter)'
TOTAL_ELIGIBLE_NOTICE = 'Total Eligible voting members: '

# create dictionary for easier access/modification
HEADINGS = { 'VOTES' : VOTES_HEADING, 'EVALUATIONS' : EVALUATIONS_HEADING }
NOTICES = { 'COMMENTS' : COMMENTS_NOTICE , 'ELIGIBLE' : TOTAL_ELIGIBLE_NOTICE }

# open document, this document contains a custom borderless table style called 'BorderlessTable'
doc = Document('template.docx')

# create table for displaying result heading
headingTable = doc.add_table(rows=2, cols=2)

# change table style to custom table style that hides border lines
headingTable.style = 'borderlessTable'
setColumnWidth(headingTable.columns[COL_ONE],Inches(1))
setColumnWidth(headingTable.columns[COL_TWO],Inches(5.5))

# add date information to first row of the heading table
headingTable.cell(0,0).text = 'DATE:'
headingTable.cell(0,1).text = getTodaysDate()

# add description second row
try:
	# get description from poll data
	pollData = getPollData(conn,pollId)
	descriptions = getActionDescriptions(pollData)

	# display discription
	headingTable.cell(1,0).text = 'RE:'
	headingDescription = descriptions[0]
	headingTable.cell(1,1).text = headingDescription

except:
	print "Could not retreive poll description"

# # format horizontal line (divider)
# doc.add_paragraph('', style='horizontalLine')
# doc.add_paragraph('', style='emptyLine')

# poll variables
# pollData and descriptions set on line 80
voteCounts = getVoteCounts(conn,pollData)
pollType = pollData['pollType']

# # printing variables
actionNum = eligible = 0

for description in descriptions:
	# begin from first action, subtract 1 from actionNum for indexing
	actionNum += 1
	romanActionNum = intToRoman(actionNum)

	# print action vote tally before printing action evaluations (comments)
	printActionVoteResults = TRUE
	printActionEvaluations = FALSE

	# print both types of action results/information
	while(printActionVoteResults or printActionEvaluations):
		# reset variables
		eligible = 0

		# start with vote tally
		if(printActionVoteResults):
			# set action result variables
			if(isMultiActionPollType(pollType)):
				eligible = voteCounts[str(actionNum)]['eligible']
			else:
				eligible = voteCounts['eligible']

			actionResults = formatActionResults(conn,pollData,actionNum)
			eligibleSummary = NOTICES['ELIGIBLE'] + str(eligible)

			# # create table to format results
			# actionTable = doc.add_table(rows=5,cols=2)
			# actionTable.style = 'borderlessTable'
			# setColumnWidth(actionTable.columns[COL_ONE],Inches(0.5))
			# setColumnWidth(actionTable.columns[COL_TWO],Inches(6))
			# # display action num and votes heading
			# actionTable.cell(0,0).text = romanActionNum
			# # print heading
			# actionTable.cell(0,1).text = HEADINGS['VOTES']
			# # print action description
			# actionTable.cell(1,1).text = description
			# # print total eligible notice
			# actionTable.cell(2,1).text = eligibleSummary
			# # print action results
			# actionTable.cell(3,1).text = actionResults
			# # print comments notice
			# actionTable.cell(4,1).text = NOTICES['COMMENTS']

			# begin printing action comments, if any
			printComments(conn,pollData,actionNum)

			# # format horizontal line (divider)
			# doc.add_paragraph('', style='horizontalLine')
			# doc.add_paragraph('', style='emptyLine')

			# end of printing vote tally and comments, start printing evaluations
			printActionVoteResults = FALSE
			printActionEvaluations = TRUE

		elif(printActionEvaluations):
		# 	# there may not be evaluations for this action
		# 	evaluations = checkForEvaluations(conn,pollData,actionNum)

		# 	# print evaluations if there are comments
		# 	if(evaluations['commentCount'] > 0):
		# 		# evaluation variables
		# 		eligibleSummary = NOTICES['ELIGIBLE'] + str(evaluations['eligible'])

		# 		# create table to format evaluations
		# 		evaluationsTable = doc.add_table(rows=4,cols=2)
		# 		evaluationsTable.style = 'borderlessTable'
		# 		setColumnWidth(evaluationsTable.columns[COL_ONE],Inches(0.5))
		# 		setColumnWidth(evaluationsTable.columns[COL_TWO],Inches(6))
				
		# 		# display action num and votes heading
		# 		evaluationsTable.cell(0,0).text = romanActionNum + 'a'
		# 		# print heading
		# 		evaluationsTable.cell(0,1).text = HEADINGS['EVALUATIONS']
		# 		# print action description
		# 		evaluationsTable.cell(1,1).text = description
		# 		# print total eligible notice
		# 		evaluationsTable.cell(2,1).text = eligibleSummary
		# 		# print comments notice
		# 		evaluationsTable.cell(3,1).text = NOTICES['COMMENTS']
		# 		# print comments
		# 		# printComments(conn,pollData,actionNum,IS_EVALUATION)


		# 		# format horizontal line (divider)
		# 		doc.add_paragraph('', style='horizontalLine')
		# 		doc.add_paragraph('', style='emptyLine')

			# begin printing action comments, if any
			# end of printing evaluations
			printActionEvaluations = FALSE

#close database connection
conn.close()

# # store document with another name to prevent losing the custom table style
# doc.save('test.docx')